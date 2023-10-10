import customtkinter
from PIL import Image
import os
import copy
import csv
from docx import Document # for writing the word document output
from docx.shared import Inches # for writing the word document output
from tkinter import messagebox
from tkinter import filedialog


customtkinter.set_appearance_mode("light")
customtkinter.set_default_color_theme("dark-blue")

class App(customtkinter.CTk): # Creating a class for the app
    def __init__(self): # initialisation function
        super().__init__()
        
        self.title("Mobile Duress Suitability Software")

        self.WIDTH = 1920
        self.HEIGHT = 1080

        #----PAGE DECLARATIONS----#

        self.geometry(f"{self.WIDTH}x{self.HEIGHT}") # setting the window size

        ##----VARIABLES FOR STORING QUESTION LOGIC----#
        self.reset_variables()

        customtkinter.set_default_color_theme("style/styles_debranded.json")
        self.char_dict = self.build_characteristics_dict()
        # Build the dictionaries for getting the text that sits on the question pages and the results page.
        self.question_topic_dict = self.build_question_topic_dict()
        self.question_text_dict = self.build_question_text_dict()
        self.question_desc_dict = self.build_question_desc_dict()
        self.tech_desc_dict = self.build_tech_desc_dict()
        self.gen_desc_dict = self.build_general_text_dict()
        self.build_welcome_page()
        #self.build_results_page() # Just for testing purposes

    def reset_variables(self):
        self.current_question = "0" # Stores the string that indicates which question that the user is currently on ("0" refers to the welcome page)
        self.question_history = [] # array that stores the history of questions that the user has been to
        self.choice_history = [] # array that stores the history of the user's choice at each question, this will be used to generate the characteristics displayed on the results page and the scoring performed to generate the top choice.
        self.scores = { # Python dictionary to keep track of the scores
            "BLE": 0,
            "Wi-Fi BLE Hybrid": 0,
            "Wi-Fi": 0,
            "RF": 0,
            "DECT": 0
        }
        self.good_char = { # Python dictionary to store all of the good characteristics that each technology has based on the user's choices
            "BLE": [],
            "Wi-Fi BLE Hybrid": [],
            "Wi-Fi": [],
            "RF": [],
            "DECT": []
        }
        self.bad_char = { # Python dictionary to store all of the bad characterstics that each technology has based on the user's choices
            "BLE": [],
            "Wi-Fi BLE Hybrid": [],
            "Wi-Fi": [],
            "RF": [],
            "DECT": []
        }

    def build_characteristics_dict(self): # Function to read the external file and build the database of characteristics in a python dictionary
        filename = "char_src.csv" # Define the filename of the source file of the dictionary
        char_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    char_dict[int(row[0])] = row[1] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to the 'statement ID'
        return char_dict

    def build_question_topic_dict(self): # Function to read the external file and build a database of the text that is used for the questions topic name
        filename = "question_src.csv" # Define the filename of the source file of the dictionary
        question_topic_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    question_topic_dict[row[0]] = row[1] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to the 'Question ID'
        return question_topic_dict
    
    def build_question_text_dict(self): # Function to read the external file and build a database of the text that is used for the questions themselves
        filename = "question_src.csv" # Define the filename of the source file of the dictionary
        question_text_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    question_text_dict[row[0]] = row[2] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to the 'Question ID'
        return question_text_dict
    
    def build_question_desc_dict(self): # Function to read the external file and build a database of the text that is used for the text that describes the question topics
        filename = "question_src.csv" # Define the filename of the source file of the dictionary
        question_desc_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    question_desc_dict[row[0]] = row[3] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to the 'Question ID'
        return question_desc_dict
    
    def build_tech_desc_dict(self): # Function to read the external file and build a database of the text that is used for the descriptions of the technologies on the results page
        filename = "techdesc_src.csv" # Define the filename of the source file of the dictionary
        tech_desc_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    tech_desc_dict[row[0]] = row[1] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to the technology the description describes
        return tech_desc_dict
    
    def build_general_text_dict(self): # Function to read the external file and build a database of the text that will be used throughout the results page and the report. Used for longer more generic blocks of text that largely do not depend on the result of using the software.
        filename = "gendesc_src.csv" # Define the filename of the source file of the dictionary
        gen_desc_dict = {} # Characteristic dictionary initialisation
        with open(filename, 'r') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if len(row) >= 2: # Validate that the file has at least 2 columns
                    gen_desc_dict[row[0]] = row[1] # make a new key-value pair in the database with the second col of the CSV file. The key is the first col which correponds to what the text is describing
        return gen_desc_dict

    def build_welcome_page(self):
        # Initialising variables that will be used throughout the program logic
        self.reset_variables()

        try:
            self.QuestionPage.destroy()
        except AttributeError:
            print("Question Page does not exist")
        try:
            self.ResultsPage.destroy()
        except AttributeError:
            print("Results Page does not exist")

        #----GRID CONFIGURAIONS----
        self.WelcomePage = customtkinter.CTkFrame(self, fg_color="white", width=self.WIDTH, height=self.HEIGHT)
        self.WelcomePage.pack(fill=customtkinter.BOTH)
        # Adjusting the weights of the rows in the window in the case of resizing
        self.WelcomePage.grid_rowconfigure((0,1), weight=1)
        self.WelcomePage.grid_rowconfigure(6, weight=10)
        self.WelcomePage.grid_rowconfigure((3,4,5), weight=2)
        self.WelcomePage.grid_columnconfigure((1,2), weight=5)
        self.WelcomePage.grid_columnconfigure((0,3), weight=1)

        #----WIDGET CONFIGURATIONS----#
        # Top bar
        self.WelcomePage.logo_frame = customtkinter.CTkFrame(self.WelcomePage, fg_color='#183E6F', height=60)
        self.WelcomePage.logo_frame.grid(row=0, column=0, columnspan=5, rowspan=1, sticky="new")
        self.WelcomePage.company_logo_label = customtkinter.CTkLabel(self.WelcomePage.logo_frame, text="Mobile Duress Suitability Software", font=customtkinter.CTkFont(size=20, weight="normal"), text_color="white", fg_color="#183E6F", width=120)  # display image with a CTkLabel
        self.WelcomePage.company_logo_label.grid(row=0, column=0, padx=10, pady=15)

        # Create a frame for each of the buttons and blank spaces in the toolbar
        # First frame lies underneath to set the background colour
        self.WelcomePage.buttonbar_frame = customtkinter.CTkFrame(self.WelcomePage,width=self.WIDTH, height=60, fg_color='#F05A5F')
        self.WelcomePage.buttonbar_frame.grid(row=1,column=0, columnspan=5, sticky="nesw")
        #Following grid configurations follow the same as the main page
        self.WelcomePage.buttonbar_frame.grid_columnconfigure((1,2,3), weight=5)
        self.WelcomePage.buttonbar_frame.grid_columnconfigure((0,4), weight=1)

        # Following frames separate the buttons into their own frames to align their spacing
        self.WelcomePage.buttonbar_frame0 = customtkinter.CTkFrame(self.WelcomePage.buttonbar_frame, width=200, height=60, fg_color='#F05A5F')
        self.WelcomePage.buttonbar_frame0.grid(row=1,column=0, columnspan=1, sticky="nesw")
        self.WelcomePage.buttonbar_frame1 = customtkinter.CTkFrame(self.WelcomePage.buttonbar_frame, width=100, height=60, fg_color='#F05A5F')
        self.WelcomePage.buttonbar_frame1.grid(row=1,column=1, columnspan=1, sticky="nsew") 
        self.WelcomePage.buttonbar_frame2 = customtkinter.CTkFrame(self.WelcomePage.buttonbar_frame, width=100, height=60, fg_color='#F05A5F')
        self.WelcomePage.buttonbar_frame2.grid(row=1,column=2, columnspan=1, sticky="nesw") 
        self.WelcomePage.buttonbar_frame3 = customtkinter.CTkFrame(self.WelcomePage.buttonbar_frame, width=100, height=60, fg_color='#F05A5F')
        self.WelcomePage.buttonbar_frame3.grid(row=1,column=3, columnspan=1, sticky="nesw")
        self.WelcomePage.buttonbar_frame4 = customtkinter.CTkFrame(self.WelcomePage.buttonbar_frame, width=200, height=60, fg_color='#F05A5F',)
        self.WelcomePage.buttonbar_frame4.grid(row=1,column=4, columnspan=1, sticky="nesw") 

        # Creating main content frames using grid naming convention
        # Top row of white content area
        self.WelcomePage.content_frame00 = customtkinter.CTkFrame(self.WelcomePage, width=self.WIDTH, height=100, fg_color="#ffffff", border_color="#ffffff") 
        self.WelcomePage.content_frame00.grid(row=2, column=0, columnspan=4, sticky="nesw")

        # First row of white content
        self.WelcomePage.content_frame12 = customtkinter.CTkFrame(self.WelcomePage, fg_color="#ffffff", border_color="#ffffff") # frame for welcome label
        self.WelcomePage.content_frame12.grid(row=3, column=0, columnspan=4, sticky="nsew")
        
        # Second row of white content
        self.WelcomePage.content_frame20 = customtkinter.CTkFrame(self.WelcomePage, width=300, height=100, fg_color="#ffffff", border_color="#ffffff") 
        self.WelcomePage.content_frame20.grid(row=4, column=0, columnspan=1, sticky="nesw")
        self.WelcomePage.content_frame2 = customtkinter.CTkFrame(self.WelcomePage, width=200, height=100, fg_color="#ffffff", border_color="#ffffff") # frame for guidance text
        self.WelcomePage.content_frame2.grid(row=4, column=1, columnspan=2, sticky="nesw")
        self.WelcomePage.content_frame23 = customtkinter.CTkFrame(self.WelcomePage, width=300, height=100, fg_color="#ffffff", border_color="#ffffff") 
        self.WelcomePage.content_frame23.grid(row=4, column=3, columnspan=1, sticky="nesw")

        # Third row of white content
        self.WelcomePage.content_frame30 = customtkinter.CTkFrame(self.WelcomePage, width=150, height=50, fg_color="#ffffff", border_color="#ffffff") 
        self.WelcomePage.content_frame30.grid(row=5, column=0, columnspan=1, sticky="nesw")
        self.WelcomePage.content_frame31 = customtkinter.CTkFrame(self.WelcomePage, fg_color='#ffffff', border_color="#ffffff") # frame for load button
        self.WelcomePage.content_frame31.grid(row=5,column=1, columnspan=1, sticky="nsew") 
        self.WelcomePage.content_frame32 = customtkinter.CTkFrame(self.WelcomePage, fg_color='#ffffff', border_color="#ffffff")# frame for new project button
        self.WelcomePage.content_frame32.grid(row=5, column=2, columnspan=1, sticky="nsew") 
        # self.WelcomePage.content_frame33 = customtkinter.CTkFrame(self.WelcomePage, fg_color='#ffffff', border_color="#ffffff") # frame for home button
        # self.WelcomePage.content_frame33.grid(row=5, column=3, columnspan=1, sticky="nsew")
        self.WelcomePage.content_frame34 = customtkinter.CTkFrame(self.WelcomePage, width=150, height=50, fg_color='#ffffff', border_color="#ffffff")
        self.WelcomePage.content_frame34.grid(row=5, column=3, columnspan=1, sticky="nsew") 

        # Bottom row of white content area
        self.WelcomePage.content_frame40 = customtkinter.CTkFrame(self.WelcomePage, width=self.WIDTH, fg_color="white") 
        self.WelcomePage.content_frame40.grid(row=6, column=0, columnspan=4, sticky="nesw")

        #---LABELS & BUTTONS---#
        # Welcome label
        self.WelcomePage.welcome_label = customtkinter.CTkLabel(self.WelcomePage.content_frame12,text="Welcome!", font=customtkinter.CTkFont(size=56, weight="bold"), text_color='#183E6F', fg_color="transparent", anchor="center")
        self.WelcomePage.welcome_label.pack()

        # Guidance text
        self.WelcomePage.guidance_label = customtkinter.CTkLabel(self.WelcomePage.content_frame2, text="Click one of the buttons below to get started.", font=customtkinter.CTkFont(size=14), text_color='#183E6F', width=600)
        self.WelcomePage.guidance_label.pack()

        # Creating buttons
        self.WelcomePage.load_button_top = customtkinter.CTkButton(self.WelcomePage.buttonbar_frame1, command=self.load_project, text='Load Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.WelcomePage.load_button_top.pack(padx=10, pady=10)
        self.WelcomePage.new_project_button_top = customtkinter.CTkButton(self.WelcomePage.buttonbar_frame3, command=self.changeToQuestion1, text='New Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.WelcomePage.new_project_button_top.pack(padx=10, pady=10)
        self.WelcomePage.home_button_top = customtkinter.CTkButton(self.WelcomePage.buttonbar_frame2, command=self.home_button_event, text='Home', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.WelcomePage.home_button_top.pack(padx=10, pady=10)

        self.WelcomePage.load_button_mid = customtkinter.CTkButton(self.WelcomePage.content_frame31, command=self.load_project, text='Load Project', fg_color="#F05A5F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#ffcb77")
        self.WelcomePage.load_button_mid.pack()
        self.WelcomePage.new_project_button_mid = customtkinter.CTkButton(self.WelcomePage.content_frame32, command=self.changeToQuestion1, text='New Project', fg_color="#F05A5F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#ffcb77")
        self.WelcomePage.new_project_button_mid.pack()
        # self.WelcomePage.home_button_mid = customtkinter.CTkButton(self.WelcomePage.content_frame32, command=self.home_button_event, text='User Manual', fg_color="#F05A5F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        # self.WelcomePage.home_button_mid.pack()
        self.WelcomePage.author_label = customtkinter.CTkLabel(self.WelcomePage.content_frame40, text="Created by Kevin Sein Tu, 2023", font=customtkinter.CTkFont(size=8), text_color='grey', height=200)
        self.WelcomePage.author_label.pack(side="bottom")
        #-----------------------------------------------------------------#

    def changeToQuestion1(self): # Function to change from welcome page and question 1
        ##----------QUESTION 1 PAGE----------##
        print("changeToQuestion1")
        try:
            self.WelcomePage.destroy()
        except AttributeError:
            print("Welcome Page does not exist")
        try:
            self.ResultsPage.destroy()
        except AttributeError:
            print("Results Page does not exist")

        self.current_question = "1" # Now moved onto the first question
        self.question_history.append("1")

        self.QuestionPage = customtkinter.CTkFrame(self, fg_color="white", width=self.WIDTH, height=self.HEIGHT) # Creating frame to span the whole page
        self.QuestionPage.grid_rowconfigure((0,1), weight=1)
        self.QuestionPage.grid_rowconfigure((2,3,4), weight=5)
        self.QuestionPage.grid_rowconfigure(5, weight=6)
        self.QuestionPage.grid_columnconfigure(0, weight=1)
        self.QuestionPage.grid_columnconfigure((1,2,3,4), weight=8)

        #----WIDGET CONFIGURATIONS----#
        # Top bar
        self.QuestionPage.logo_frame = customtkinter.CTkFrame(self.QuestionPage, fg_color='#183E6F', border_color='#183E6F')
        self.QuestionPage.logo_frame.grid(row=0, column=0, columnspan=5, rowspan=1, sticky="new")
        self.QuestionPage.company_logo_label = customtkinter.CTkLabel(self.QuestionPage.logo_frame, text="Mobile Duress Suitability Software", font=customtkinter.CTkFont(size=20, weight="normal"), text_color="white", fg_color="#183E6F", width=120)  # display image with a CTkLabel
        self.QuestionPage.company_logo_label.grid(row=0, column=0, padx=10, pady=15)

        # Create a frame for each of the buttons and blank spaces in the toolbar
        # First frame lies underneath to set the background colour
        self.QuestionPage.buttonbar_frame = customtkinter.CTkFrame(self.QuestionPage,width=self.WIDTH-200, height=60, fg_color='#F05A5F')
        self.QuestionPage.buttonbar_frame.grid(row=1,column=0, columnspan=5, sticky="nesw")
        #Following grid configurations follow the same as the main page
        self.QuestionPage.buttonbar_frame.grid_columnconfigure((1,2,3), weight=3)
        self.QuestionPage.buttonbar_frame.grid_columnconfigure(0, weight=2)
        self.QuestionPage.buttonbar_frame.grid_columnconfigure(4, weight=2)

        # Following frames separate the buttons into their own frames to align their spacing
        self.QuestionPage.buttonbar_frame0 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=100, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.QuestionPage.buttonbar_frame0.grid(row=1,column=0, columnspan=1, sticky="nesw")
        self.QuestionPage.buttonbar_frame1 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.QuestionPage.buttonbar_frame1.grid(row=1,column=1, columnspan=1, sticky="nsew") 
        self.QuestionPage.buttonbar_frame2 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.QuestionPage.buttonbar_frame2.grid(row=1,column=2, columnspan=1, sticky="nesw") 
        self.QuestionPage.buttonbar_frame3 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.QuestionPage.buttonbar_frame3.grid(row=1,column=3, columnspan=1, sticky="nesw")
        self.QuestionPage.buttonbar_frame4 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=60, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.QuestionPage.buttonbar_frame4.grid(row=1,column=4, columnspan=1, sticky="nesw")
        # self.QuestionPage.buttonbar_frame5 = customtkinter.CTkFrame(self.QuestionPage.buttonbar_frame, width=60, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        # self.QuestionPage.buttonbar_frame5.grid(row=1,column=5, columnspan=1, sticky="nesw") 
        
        # Creating a frame to hold all of the content frames in the white space
        self.content_height=self.HEIGHT-self.QuestionPage.buttonbar_frame.cget("height")-self.QuestionPage.logo_frame.cget("height")-100
        self.QuestionPage.content_frame = customtkinter.CTkFrame(self.QuestionPage, fg_color="white", width=self.WIDTH, height=self.content_height) 
        self.QuestionPage.content_frame.grid(row=2, column=1, rowspan=4, columnspan=4, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.content_frame.grid_rowconfigure(0, weight=2)
        self.QuestionPage.content_frame.grid_rowconfigure((1,2), weight=1)
        self.QuestionPage.content_frame.grid_rowconfigure(3, weight=6)
        self.QuestionPage.content_frame.grid_rowconfigure(4, weight=1)
        self.QuestionPage.content_frame.grid_columnconfigure(0, weight=2)

        # Creating buttons
        # self.QuestionPage.save_project_button = customtkinter.CTkButton(self.QuestionPage.buttonbar_frame1, command=self.home_button_event, text='Save Current Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        # self.QuestionPage.save_project_button.pack(padx=5, pady=10)
        self.QuestionPage.load_button = customtkinter.CTkButton(self.QuestionPage.buttonbar_frame1, command=self.load_project, text='Load Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.QuestionPage.load_button.pack(padx=5, pady=10)
        self.QuestionPage.new_project_button = customtkinter.CTkButton(self.QuestionPage.buttonbar_frame3, command=self.changeToQuestion1, text='New Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.QuestionPage.new_project_button.pack(padx=5, pady=10)
        self.QuestionPage.home_button = customtkinter.CTkButton(self.QuestionPage.buttonbar_frame2, command=self.home_button_event, text='Home', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.QuestionPage.home_button.pack(padx=5, pady=10)

        # Creating text labels in main body
        self.QuestionPage.question_title = customtkinter.CTkLabel(self.QuestionPage.content_frame, text=self.get_question_text("1")[0], font=customtkinter.CTkFont(size=24, weight="bold"), text_color="#183E6F", anchor="w", wraplength=1200, justify="left")
        self.QuestionPage.question_title.grid(row=0, column=0, columnspan=1, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.question_desc = customtkinter.CTkLabel(self.QuestionPage.content_frame, text=self.get_question_text("1")[1], font=customtkinter.CTkFont(size=16, weight="normal"), text_color="#183E6F", anchor="w", wraplength=1350, justify="left")
        self.QuestionPage.question_desc.grid(row=1, column=0, columnspan=1, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.question_text = customtkinter.CTkLabel(self.QuestionPage.content_frame, text=self.get_question_text("1")[2], font=customtkinter.CTkFont(size=16, weight="bold"), text_color="#183E6F", anchor="w", wraplength=1200, justify="left")
        self.QuestionPage.question_text.grid(row=2, column=0, columnspan=1, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.option_frame = customtkinter.CTkFrame(self.QuestionPage.content_frame, width=100, height=200)
        self.QuestionPage.option_frame.grid(row=3, column=0, columnspan=1, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.prevnext_frame = customtkinter.CTkFrame(self.QuestionPage.content_frame)
        self.QuestionPage.prevnext_frame.grid(row=4, column=0, columnspan=1, sticky="nsew", padx=20, pady=10)
        self.QuestionPage.prevnext_frame.grid_columnconfigure(0,weight=15)
        self.QuestionPage.prevnext_frame.grid_columnconfigure((1,2),weight=1)
        self.QuestionPage.prevnext_frame.grid_rowconfigure(0,weight=1)
        
        # Generating the radio button options
        self.generate_radio_buttons("1") # Generate the radio buttons for the first question

        # Creating Previous and next buttons at the bottom of the page
        self.QuestionPage.prev_button = customtkinter.CTkButton(self.QuestionPage.prevnext_frame, command=self.callback_prev_button, text='Previous', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.QuestionPage.prev_button.grid(row=0, column=1)
        self.QuestionPage.prev_button.configure(state="disabled") # Disable the previous question button since we are on the first question
        self.QuestionPage.next_button = customtkinter.CTkButton(self.QuestionPage.prevnext_frame, command=self.callback_next_button, text='Next', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#a0ddff")
        self.QuestionPage.next_button.grid(row=0, column=2)
        self.QuestionPage.next_button.configure(state="disabled") # Can't go to the next question until we put in a radio option

        self.QuestionPage.pack(fill=customtkinter.BOTH)

# Set of callback functions to abstract the execution of commands that are done when different buttons in the program are clicked
    def cond_next_button(self):
        # If we are not on the last question and we have selected a radio option
        #print(self.current_question, self.QuestionPage.radio_option.get())
       # print(self.current_question != "9e" and self.QuestionPage.radio_option.get()!=0)
        #return (self.current_question != "9e" and self.QuestionPage.radio_option.get()!=0)
        return (self.QuestionPage.radio_option.get()!=0)
    def cond_prev_button(self):
        return (self.current_question != "1")

    def callback_next_button(self):
        #self.increase_scores(self.current_question, self.QuestionPage.radio_option.get()) # adjust the scores
        # Change the current question we are on
        self.changeQuestion(self.current_question, 1)
        
        # Enable/disable the next button depending on which question we end up on
        if self.cond_next_button(): # If we satisfy the conditions to go to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")
        if self.cond_prev_button(): # If we satisfy the conditions to go back to the previous question
            self.QuestionPage.prev_button.configure(state="enabled")
        else:
            self.QuestionPage.prev_button.configure(state="disabled")

    def callback_prev_button(self):
        self.changeQuestion(self.current_question, -1)
        # Enable/disable the next button depending on which question we end up on
        if self.cond_next_button: # If we satisfy the conditions to go to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")
        if self.cond_prev_button(): # # If we satisfy the conditions to go back to the previous question
            self.QuestionPage.prev_button.configure(state="enabled")
        else: 
            self.QuestionPage.prev_button.configure(state="disabled")
    
    def radio_option1_callback(self):
        print("Selected option:",self.QuestionPage.radio_option.get())
        # Code to change the colour framing
        self.QuestionPage.option1_frame.configure(fg_color="#a0ddff")
        self.QuestionPage.option1_radio.configure(text_color="#183E6F")
        # Change the colours of the other radio buttons back
        num_options = self.get_question_options(self.current_question)[0] # Get how many options there are for the current question
        if num_options >=2:
            self.QuestionPage.option2_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option2_radio.configure(text_color="white")
        if num_options >=3:
            self.QuestionPage.option3_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option3_radio.configure(text_color="white")
        if num_options >=4:
            self.QuestionPage.option4_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option4_radio.configure(text_color="white")
        if num_options >=5:
            self.QuestionPage.option5_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option5_radio.configure(text_color="white")
        if self.cond_next_button(): # # If we satisfy the conditions to go back to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")
    
    def radio_option2_callback(self):
        print("Selected option:",self.QuestionPage.radio_option.get())
        # Code to change the colour framing
        self.QuestionPage.option2_frame.configure(fg_color="#a0ddff")
        self.QuestionPage.option2_radio.configure(text_color="#183E6F")
        # Change the colours of the other radio buttons back
        num_options = self.get_question_options(self.current_question)[0] # Get how many options there are for the current question
        if num_options >=2:
            self.QuestionPage.option1_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option1_radio.configure(text_color="white")
        if num_options >=3:
            self.QuestionPage.option3_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option3_radio.configure(text_color="white")
        if num_options >=4:
            self.QuestionPage.option4_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option4_radio.configure(text_color="white")
        if num_options >=5:
            self.QuestionPage.option5_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option5_radio.configure(text_color="white")
        if self.cond_next_button(): # # If we satisfy the conditions to go back to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")

    def radio_option3_callback(self):
        print("Selected option:",self.QuestionPage.radio_option.get())
        # Code to change the colour framing
        self.QuestionPage.option3_frame.configure(fg_color="#a0ddff")
        self.QuestionPage.option3_radio.configure(text_color="#183E6F")
        # Change the colours of the other radio buttons back
        num_options = self.get_question_options(self.current_question)[0] # Get how many options there are for the current question
        if num_options >=3:
            self.QuestionPage.option1_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option1_radio.configure(text_color="white")
            self.QuestionPage.option2_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option2_radio.configure(text_color="white")
        if num_options >=4:
            self.QuestionPage.option4_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option4_radio.configure(text_color="white")
        if num_options >=5:
            self.QuestionPage.option5_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option5_radio.configure(text_color="white")
        if self.cond_next_button(): # # If we satisfy the conditions to go back to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")
    
    def radio_option4_callback(self):
        print("Selected option:",self.QuestionPage.radio_option.get())
        # Code to change the colour framing
        self.QuestionPage.option4_frame.configure(fg_color="#a0ddff")
        self.QuestionPage.option4_radio.configure(text_color="#183E6F")
        # Change the colours of the other buttons back
        num_options = self.get_question_options(self.current_question)[0]
        self.QuestionPage.option1_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option1_radio.configure(text_color="white")
        self.QuestionPage.option2_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option2_radio.configure(text_color="white")
        self.QuestionPage.option3_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option3_radio.configure(text_color="white")
        if num_options >=5:
            self.QuestionPage.option5_frame.configure(fg_color="#183E6F")
            self.QuestionPage.option5_radio.configure(text_color="white")
        if self.cond_next_button(): # # If we satisfy the conditions to go back to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")
        
    def radio_option5_callback(self):
        print("Selected option:",self.QuestionPage.radio_option.get())
        # Code to change the colour framing
        self.QuestionPage.option5_frame.configure(fg_color="#a0ddff")
        self.QuestionPage.option5_radio.configure(text_color="#183E6F")
        # Change the colours of the other buttons back
        num_options = self.get_question_options(self.current_question)[0]
        self.QuestionPage.option1_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option1_radio.configure(text_color="white")
        self.QuestionPage.option2_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option2_radio.configure(text_color="white")
        self.QuestionPage.option3_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option3_radio.configure(text_color="white")
        self.QuestionPage.option4_frame.configure(fg_color="#183E6F")
        self.QuestionPage.option4_radio.configure(text_color="white")
    
        if self.cond_next_button(): # # If we satisfy the conditions to go back to the next question
            self.QuestionPage.next_button.configure(state="enabled")
        else:
            self.QuestionPage.next_button.configure(state="disabled")

    def get_question_text(self, question_index): # Function to operate as a lookup table for which question to get the text for, returns a list (question_topic, topic_desc, question_text, question type)
        question_text = self.question_text_dict[question_index]
        question_desc = self.question_desc_dict[question_index]
        question_topic = self.question_topic_dict[question_index]

        return question_topic, question_desc, question_text
        
    def get_question_options(self, question_index): # Function to operate as a lookup table for options to provide to the user (option1, option2, ...)
        match question_index: #Switch case statement to get the text for each of the options within the question
            case "1": # Each case is one of the questions in the decision tree
                question_option1 = "Within 3m"
                question_option2 = "Between 3m and 5m"
                question_option3 = "Greater than 5m"
                num_options = 3
            case "2": # Each case is one of the questions in the decision tree
                question_option1 = "New Facility"
                question_option2 = "Existing Facility"
                num_options = 2            
            case "3a": # Each case is one of the questions in the decision tree
                question_option1 = "Yes"
                question_option2 = "No"
                num_options = 2         
            case "3b": # Each case is one of the questions in the decision tree
                question_option1 = "Voice-grade or Data-grade"
                question_option2 = "RTLS-grade"
                num_options = 2         
            case "4": # Each case is one of the questions in the decision tree
                question_option1 = "Low-medium"
                question_option2 = "High"
                num_options = 2         
            case "5a": # Each case is one of the questions in the decision tree
                question_option1 = "Low-medium"
                question_option2 = "High"
                num_options = 2  
            case "5b": # Each case is one of the questions in the decision tree
                question_option1 = "Low-medium"
                question_option2 = "High"
                num_options = 2  
            case "6": # Each case is one of the questions in the decision tree
                question_option1 = "Yes"
                question_option2 = "No"
                num_options = 2        
            case "7": # Each case is one of the questions in the decision tree
                question_option1 = "Smartphone, Wi-Fi and bluetooth capable"
                question_option2 = "Smartphone, Wi-Fi capable only"
                question_option3 = "DECT Handset or pager"
                question_option4 = "RF Handset or pager"
                question_option5 = "Locator Tag"
                num_options = 5        
            # case "8": # Each case is one of the questions in the decision tree
            #     question_option1 = "Smartphone, Wi-Fi and bluetooth capable"
            #     question_option2 = "Smartphone, Wi-Fi capable only"
            #     question_option3 = "DECT Handset"
            #     question_option4 = "RF Handset"
            #     num_options = 4          
            case "9a": # Each case is one of the questions in the decision tree
                question_option1 = "Thick walls"
                question_option2 = "Thin walls or medium walls or not sure"
                num_options = 2   
            # case "9b": # Each case is one of the questions in the decision tree
            #     question_option1 = "Yes"
            #     question_option2 = "No"
            #     num_options = 2   
            case "9c": # Each case is one of the questions in the decision tree
                question_option1 = "Yes"
                question_option2 = "No"
                num_options = 2   
            case "9d": # Each case is one of the questions in the decision tree
                question_option1 = "Yes"
                question_option2 = "No"
                num_options = 2   
            case "9e": # Each case is one of the questions in the decision tree
                question_option1 = "Yes"
                question_option2 = "No"
                num_options = 2   
            case _:
                question_option1 = "<EXCEPTION REACHED>"
                question_option2 = "<EXCEPTION REACHED>"
                num_options=2

        match num_options: # Return the options based on whether or not they exist.
            case 2: # If the question has 2 selectable options
                return  num_options, question_option1, question_option2
            case 3: # If the question has 3 selectable options
                return num_options, question_option1, question_option2, question_option3
            case 4: # If the question has 4 selectable options
                return num_options, question_option1, question_option2, question_option3, question_option4
            case 5: # If the question has 4 selectable options
                return num_options, question_option1, question_option2, question_option3, question_option4, question_option5
            case _: # exception case
                print("<EXCEPTION REACHED>")

    def generate_radio_buttons(self, question_index): # Function for dynamically generating the radio buttons that sit on the bottom of the question page
        # Check if this is the first function call
        print("Current_question:" + self.current_question)
        #if (self.question_history[-1] != "1" or (self.question_history[-1] == "1" and self.question_history[-2] == "2")): # If we are not on question 1 or going back to question 1
            # The option frames will already exist in all of these cases, since the options are first generated when going from the welcome page to question 1
        try:
            self.QuestionPage.option1_frame.destroy()
        except AttributeError:
            print("Option1 doesn't exist")
        try:
            self.QuestionPage.option2_frame.destroy()
        except AttributeError:
            print("Option2 doesn't exist")
        try:
            self.QuestionPage.option3_frame.destroy()
        except AttributeError:
            print("Option3 doesn't exist")
        try:
            self.QuestionPage.option4_frame.destroy()
        except AttributeError:
            print("Option4 doesn't exist")
        try:
            self.QuestionPage.option5_frame.destroy()
        except AttributeError:
            print("Option5 doesn't exist")
            # The above is only assuming that the first question has 2 options, pack_forget() on more options must be added if this changes.
            # Destroy the option 3 and 4 frames if they exist
            
        # Creating radio buttons
        radio_options = self.get_question_options(question_index) # Get the options to put inside the radio buttons

        self.QuestionPage.radio_option = customtkinter.IntVar(value=0) # Create a variable for storing the value of the radio button
        self.QuestionPage.option1_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=60, corner_radius=8)
        #self.QuestionPage.option1_frame.grid(row=0, column=0, padx=10, sticky="ew")
        self.QuestionPage.option1_frame.pack(pady=10, padx=30, anchor="w")
        self.QuestionPage.option1_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option1_frame, command=self.radio_option1_callback, variable=self.QuestionPage.radio_option, value=1, font=customtkinter.CTkFont(size=16), text=radio_options[1], text_color="white", corner_radius=50)
        self.QuestionPage.option1_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
        self.QuestionPage.option1_frame.bind("<Button-1>", self.QuestionPage.option1_radio.invoke)

        self.QuestionPage.option2_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=60, corner_radius=8)
        #self.QuestionPage.option2_frame.grid(row=1, column=0, padx=10, sticky="ew")
        self.QuestionPage.option2_frame.pack(pady=10, padx=30, anchor="w")
        self.QuestionPage.option2_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option2_frame, command=self.radio_option2_callback, variable=self.QuestionPage.radio_option, value=2, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[2], text_color="white", corner_radius=50)
        self.QuestionPage.option2_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
        self.QuestionPage.option2_frame.bind("<Button-1>", self.QuestionPage.option2_radio.invoke)

        # Buttons 3 and 4 do not always appear, appearance is based on the number of options that the question has available
        match radio_options[0]:
            case 3: # the number of options that the question has is the last element in the returned array
                self.QuestionPage.option3_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option3_frame.grid(row=0, column=0, padx=10, sticky="ew")
                self.QuestionPage.option3_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option3_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option3_frame, command=self.radio_option3_callback, variable=self.QuestionPage.radio_option, value=3, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[3], text_color="white", corner_radius=50)
                self.QuestionPage.option3_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option3_frame.bind("<Button-1>", self.QuestionPage.option3_radio.invoke) # Allow the user to click the frame and still select the radio option
            case 4:
                self.QuestionPage.option3_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option3_frame.grid(row=0, column=0, padx=10, sticky="ew")
                self.QuestionPage.option3_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option3_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option3_frame, command=self.radio_option3_callback, variable=self.QuestionPage.radio_option, value=3, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[3], text_color="white", corner_radius=50)
                self.QuestionPage.option3_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option3_frame.bind("<Button-1>", self.QuestionPage.option3_radio.invoke)
                
                self.QuestionPage.option4_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option4_frame.grid(row=1, column=0, padx=10, sticky="ew")
                self.QuestionPage.option4_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option4_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option4_frame, command=self.radio_option4_callback, variable=self.QuestionPage.radio_option, value=4, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[4], text_color="white", corner_radius=50)
                self.QuestionPage.option4_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option4_frame.bind("<Button-1>", self.QuestionPage.option4_radio.invoke)
            case 5:
                self.QuestionPage.option3_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option3_frame.grid(row=0, column=0, padx=10, sticky="ew")
                self.QuestionPage.option3_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option3_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option3_frame, command=self.radio_option3_callback, variable=self.QuestionPage.radio_option, value=3, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[3], text_color="white", corner_radius=50)
                self.QuestionPage.option3_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option3_frame.bind("<Button-1>", self.QuestionPage.option3_radio.invoke)
                
                self.QuestionPage.option4_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option4_frame.grid(row=1, column=0, padx=10, sticky="ew")
                self.QuestionPage.option4_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option4_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option4_frame, command=self.radio_option4_callback, variable=self.QuestionPage.radio_option, value=4, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[4], text_color="white", corner_radius=50)
                self.QuestionPage.option4_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option4_frame.bind("<Button-1>", self.QuestionPage.option4_radio.invoke)
                
                self.QuestionPage.option5_frame = customtkinter.CTkFrame(self.QuestionPage.option_frame, fg_color="#183E6F", border_color="#183E6F", height=40, corner_radius=8)
                #self.QuestionPage.option4_frame.grid(row=1, column=0, padx=10, sticky="ew")
                self.QuestionPage.option5_frame.pack(pady=10, padx=30, anchor="w")
                self.QuestionPage.option5_radio = customtkinter.CTkRadioButton(master=self.QuestionPage.option5_frame, command=self.radio_option5_callback, variable=self.QuestionPage.radio_option, value=5, font=customtkinter.CTkFont(size=16), text=self.get_question_options(question_index)[5], text_color="white", corner_radius=50)
                self.QuestionPage.option5_radio.grid(row=0, column=0, pady=10, padx=30, sticky="nsew")
                self.QuestionPage.option5_frame.bind("<Button-1>", self.QuestionPage.option5_radio.invoke)

    def new_project_button_event(self):
        self.changeToQuestion1()

    def home_button_event(self):
        self.build_welcome_page()
    
    def get_next_question_index(self, radio_button_value, question_index): # Method takes the value of the current question and the value of the radio button to know which question should come next
        # Note that the value of the radio button has a different meaning depending on the question that it is coupled with
        # In general, 0 is the no option selected, 1 is the first option selected, 2 is the second option selected, etc.
        match question_index: #Switch case statement to get the text for each of the options within the question
            case "1": # Each case is one of the questions in the decision tree                    
                next_question_index = "2" # All responses to this question lead to the same result
            case "2":
                match radio_button_value: # Use the value of the radio button to determine which question the program is supposed to go to next
                    case 1:
                        next_question_index = "4"
                    case 2:
                        next_question_index = "3a"
            case "3a": 
                match radio_button_value: 
                    case 1:
                        next_question_index = "3b"
                    case 2:
                        next_question_index = "5a"
            case "3b": 
                match radio_button_value: 
                    case 1:
                        next_question_index = "5a"
                    case 2:
                        next_question_index = "5b"
            case "4": 
                next_question_index = "7"
            case "5a": 
                next_question_index = "6"
            case "5b": 
                next_question_index = "6"
            case "6": 
                next_question_index = "7"
            case "7": 
                next_question_index = "9a"
            # case "8": # Question removed due to redundancy
            #     next_question_index = "9a"
            case "9a": 
                next_question_index = "9c" # Removed Question 9b, so 9a redirects to 9c
            # case "9b": 
            #     next_question_index = "9c"
            case "9c":
                next_question_index = "9d"
            case "9d": 
                next_question_index = "9e"
            case "9e": 
                # Results case
                next_question_index = "Results"

            case _:
                next_question_index = "<EXCEPTION REACHED>"
        return next_question_index

    def changeQuestion(self, question_index, direction): # Method used to switch between questions in the software -- Triggered by the 'next' button or the chapter select
        # Direction takes values -1 or 1: -1 is to go to the previous question, +1 is to go to the next question
        # The next question that is selected is based on the flow of the decision tree, as documented in the flow chart
        # -- Note that destination (dest) question in this case could also refer to the previous question if we are trying to go to the previous question
        
        # By this stage we should have already validated whether the user is able to go back to a previous question or go to the next one
        if direction == 1: # Going to the next question
            print("Going to next question")
            dest_question_index = self.get_next_question_index(self.QuestionPage.radio_option.get(), question_index) # Get the index of the next question based on the user's response # gets the text associated with each of the questions
            self.choice_history.append(self.QuestionPage.radio_option.get()) # When the question changes, append the selected radio button choice to choice_history
            if dest_question_index == "Results": # If we have finished all the questions
                print("Final question finished, going to results page...")
                self.build_results_page()
            else:
                self.question_history.append(dest_question_index) # Add the new question to the question_history
        elif direction == -1: # Going to the previous question
            print("Going to previous question")
            dest_question_index = self.question_history[-2] # get second last item from question history
            self.question_history.pop() # Remove the most recent question from the question history
            self.choice_history.pop() # Remove the most recent answer choice from choice_history

        if dest_question_index != "Results":
            # Changes to the text associated with the question
            print("changeQuestion")
            self.question_arr = self.get_question_text(dest_question_index)
            self.QuestionPage.question_title.configure(text=self.question_arr[0])
            self.QuestionPage.question_desc.configure(text=self.question_arr[1])
            self.QuestionPage.question_text.configure(text=self.question_arr[2])

            # Change button states accordingly
            if (dest_question_index) == 1: # If we are on the first question
                # Disable the 'previous' button
                self.QuestionPage.prev_button.configure(state="disabled")

            self.current_question = dest_question_index  # Update the current question number
        
            # Changes the options associated with the question
            self.generate_radio_buttons(dest_question_index)
            
        print("Question History: " + str(self.question_history))
        print("Choice History: "+str(self.choice_history))
    
    def generate_scores(self, choice_history, question_history): # Method that calls increase_scores to perform the scoring and characteristic generation based on question_history and choice_history
        i = 0 # iterator to store which question we are up to
        for question in question_history: # for each question in the question history
            self.increase_scores(question, choice_history[i]) # Increase the scores according to which option was selected
            i += 1 # increment the iterate as we move to the next question

    def increase_scores(self, question_index, radio_button_value): # Increase the score of particular locating technology based on which question the user is on and what option they have selected
        print("Increase scores, question_index:",question_index, "radio_button_value", radio_button_value)
        match question_index: #Switch case statement to get the text for each of the options within the question            
            case "1":
                match radio_button_value: # Use the value of the radio button to determine which question the program is supposed to go to next
                    case 1: # Within 3m
                        self.scores["BLE"] += 1
                        self.scores["Wi-Fi BLE Hybrid"] += 1
                        self.scores["Wi-Fi"] -= 1
                        self.scores["RF"] -= 1
                        self.append_char(0) # append the characteristics lists according to the option selected
                    case 2: # 3-5m
                        self.scores["Wi-Fi"] += 1
                        self.scores["RF"] -= 1
                        self.scores["DECT"] -= 1
                        self.append_char(1)
                    case 3: # >5m
                        self.scores["RF"] += 1
                        self.scores["DECT"] += 1
                        self.append_char(2)
            case "3b": 
                match radio_button_value: 
                    case 1: # Voice-grade or data-grade
                        self.scores["BLE"] += 1
                        self.scores["RF"] += 1
                        self.scores["Wi-Fi"] -= 2
                        self.scores["Wi-Fi BLE Hybrid"] -= 2
                        self.append_char(3)
                    case 2: # RTLS-grade
                        self.scores["Wi-Fi BLE Hybrid"] += 1
                        self.scores["Wi-Fi"] += 2
                        self.append_char(4)
            case "4": 
                match radio_button_value: 
                    case 1: # Low-medium
                        self.scores["BLE"] += 3
                        self.scores["RF"] += 3
                        self.scores["DECT"] += 3
                        self.scores["Wi-Fi"] -= 1
                        self.append_char(5)
            case "5a": 
                match radio_button_value: 
                    case 1: # Low-medium
                        self.scores["BLE"] += 3
                        self.scores["RF"] += 3
                        self.scores["DECT"] += 3
                        self.scores["Wi-Fi"] -= 1
                        self.append_char(6)
            case "5b": 
                match radio_button_value: 
                    case 1: # Low-medium
                        self.scores["BLE"] += 2
                        self.scores["RF"] += 2
                        self.scores["Wi-Fi"] += 4
                        self.scores["Wi-Fi BLE Hybrid"] += 3
                        self.scores["DECT"] += 2
                        self.append_char(7)
            case "6": 
                match radio_button_value: 
                    case 1: # Yes
                        self.scores["BLE"] += 2
                        self.scores["RF"] += 1
                        self.append_char(8)
            case "7": 
                match radio_button_value: 
                    case 1: # Smartphone, Wi-Fi and bluetooth capable
                        self.scores["BLE"] += 4
                        self.scores["Wi-Fi BLE Hybrid"] += 4
                        self.scores["Wi-Fi"] += 3
                        self.append_char(9)
                    case 2: # Smartphone, Wi-Fi capable only
                        self.scores["Wi-Fi"] += 4
                        self.append_char(17)
                    case 3: # DECT
                        self.scores["DECT"] += 4
                        self.append_char(10)
                    case 4: # RF
                        self.scores["RF"] += 4
                        self.append_char(11)
                    case 5: # Locating tag
                        self.scores["DECT"] += -4 
                        self.append_char(18)
                        
            # case "8": # QUESTION REMOVED
            #    match radio_button_value: 
            #         case 1:
            #             self.scores["BLE"] += 2
            #             self.scores["Wi-Fi BLE Hybrid"] += 2
            #         case 2:
            #            self.scores["Wi-Fi"] += 2
            #         case 3:
            #            self.scores["DECT"] += 2
            #         case 4:
            #            self.scores["RF"] += 2
            case "9a": 
                match radio_button_value: 
                    case 1: # Thick walls
                        self.scores["BLE"] += 1
                        self.scores["Wi-Fi BLE Hybrid"] += -1
                        self.scores["Wi-Fi"] -= 1
                        self.scores["RF"] += 1
                        self.append_char(12)
            # case "9b": QUESTION REMOVED
            #     match radio_button_value: 
            #         case 1: # Yes
            #             self.scores["Wi-Fi"] -= 2
            #             self.scores["Wi-Fi BLE Hybrid"] -= 1
            #             self.append_char(13)
            case "9c":
                match radio_button_value: 
                    case 1: # Yes
                        self.scores["Wi-Fi"] -= 2
                        self.scores["Wi-Fi BLE Hybrid"] -= 2
                        self.append_char(14)
            case "9d":
                match radio_button_value: 
                    case 1: # Yes
                        self.scores["RF"] -= 2
                        self.append_char(15)
            case "9e": 
                match radio_button_value: 
                    case 1: # Yes
                        self.scores["Wi-Fi"] -= 2
                        self.scores["Wi-Fi BLE Hybrid"] -= 2
                        self.append_char(16)
            case _:
                print("<EXCEPTION REACHED -- scores not adjusted, invalid question index>")
        print(self.scores)
    
    def get_top_results(self, scores_dict): # Method that reads the scores and generates the top as the top recommendations
        scores_dict_temp = copy.copy(scores_dict) # Create a copy of the scores dictionary that we will be editing to find the top 3 scores
        print("Final scores:" + str(scores_dict))
        first_key = max(scores_dict_temp, key=lambda i: scores_dict[i]) # Get the dictionary key of the first score
        del scores_dict_temp[first_key] # Delete this dictionary key from the temporary dictionary (so we can find the second-highest score)
        second_key = max(scores_dict_temp, key=lambda i: scores_dict[i]) # Get the dictionary key of the second score
        del scores_dict_temp[second_key] # Delete this dictionary key from the temporary dictionary (so we can find the third-highest score)
        third_key = max(scores_dict_temp, key=lambda i: scores_dict[i]) # Get the dictionary key of the third score

        # Get the corresponding scores of all the dictionary keys that had the top 3 scores
        first_score = scores_dict[first_key]
        second_score = scores_dict[second_key]
        third_score = scores_dict[third_key]
    
        return first_key, second_key, third_key, first_score, second_score, third_score
        
    def append_char(self, case_id): # Takes an index and appends the corresponding statement into the good_characteristics of the appropriate technology
        match case_id:
            case 0: #1 -- within 3m
                self.good_char["BLE"].append(self.char_dict[0])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[1])
                self.bad_char["Wi-Fi"].append(self.char_dict[2])
                self.bad_char["RF"].append(self.char_dict[3])
                self.bad_char["DECT"].append(self.char_dict[4])
            case 1: #1 -- 3-5m
                self.good_char["BLE"].append(self.char_dict[5])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[6])
                self.good_char["Wi-Fi"].append(self.char_dict[7])
                self.bad_char["RF"].append(self.char_dict[8])
                self.good_char["DECT"].append(self.char_dict[9])
            case 2: #1 -- >5m
                self.good_char["BLE"].append(self.char_dict[10])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[11])
                self.good_char["Wi-Fi"].append(self.char_dict[12])
                self.good_char["RF"].append(self.char_dict[13])
                self.good_char["DECT"].append(self.char_dict[14])
            case 3: #3b -- Voice-grade or data-grade
                self.good_char["BLE"].append(self.char_dict[15])
                self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[16])
                self.good_char["Wi-Fi"].append(self.char_dict[17])
            case 4: #3b -- RTLS Grade
                self.bad_char["BLE"].append(self.char_dict[18])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[19])
                self.good_char["Wi-Fi"].append(self.char_dict[20])
                self.bad_char["RF"].append(self.char_dict[21])
                self.bad_char["DECT"].append(self.char_dict[22])
            case 5: #4 -- Low-medium
                self.good_char["BLE"].append(self.char_dict[23])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[24])
                self.bad_char["Wi-Fi"].append(self.char_dict[25])
                self.good_char["RF"].append(self.char_dict[26])
                self.good_char["DECT"].append(self.char_dict[27])
            case 6: #5a -- Low-medium
                self.good_char["BLE"].append(self.char_dict[28])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[29])
                self.bad_char["Wi-Fi"].append(self.char_dict[30])
                self.good_char["RF"].append(self.char_dict[31])
                self.good_char["DECT"].append(self.char_dict[32])
            case 7: #5b -- Low-medium
                self.good_char["BLE"].append(self.char_dict[33])
                #self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[34])
                self.good_char["Wi-Fi"].append(self.char_dict[35])
                self.bad_char["RF"].append(self.char_dict[36])
                self.bad_char["DECT"].append(self.char_dict[37])
            case 8: #6 -- yes
                self.good_char["BLE"].append(self.char_dict[38])
                self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[39])
                self.bad_char["Wi-Fi"].append(self.char_dict[40])
                self.good_char["RF"].append(self.char_dict[41])
                self.bad_char["DECT"].append(self.char_dict[42])
            case 9: #7 -- Smartphone, Wi-Fi and bluetooth capable
                self.good_char["BLE"].append(self.char_dict[43])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[44])
                self.good_char["Wi-Fi"].append(self.char_dict[45])
            case 17: #7 -- Smartphone, Wi-Fi capable only
                self.good_char["Wi-Fi"].append(self.char_dict[46])
                self.good_char["DECT"].append(self.char_dict[47])
            case 10: #7 -- DECT
                #self.good_char["Wi-Fi"].append(self.char_dict[46])
                self.good_char["DECT"].append(self.char_dict[47])
            case 11: #11 -- RF
                self.good_char["RF"].append(self.char_dict[48])
            case 12: #9a -- Thick walls
                self.good_char["BLE"].append(self.char_dict[49])
                self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[59])
                self.bad_char["Wi-Fi"].append(self.char_dict[50])
                self.good_char["RF"].append(self.char_dict[51])
            # case 13: #9b -- yes QUESITON REMOVED
            #     self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[52])
            #     self.bad_char["Wi-Fi"].append(self.char_dict[53])
            case 14: #9c -- yes
                self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[54])
                self.bad_char["Wi-Fi"].append(self.char_dict[55])
            case 15: #9d -- yes
                self.bad_char["RF"].append(self.char_dict[56])
            case 16: #9e -- yes
                self.bad_char["Wi-Fi"].append(self.char_dict[57])
                self.bad_char["Wi-Fi BLE Hybrid"].append(self.char_dict[58])
            case 17: # 7 -- Smartphone, Wi-Fi capable only
                self.good_char["Wi-Fi"].append(self.char_dict[46])
            case 18: # 7 -- Personal tag/pendant
                self.good_char["BLE"].append(self.char_dict[60])
                self.good_char["Wi-Fi BLE Hybrid"].append(self.char_dict[61])
                self.good_char["Wi-Fi"].append(self.char_dict[62])
                self.good_char["RF"].append(self.char_dict[63])
                self.good_char["DECT"].append(self.char_dict[64])
            case _:
                print("<EXCEPTION REACHED IN APPEND_CHAR>")

    def get_char_text(self, good_char_dict, bad_char_dict, dict_key): # To be used as the text for the label that sits on the results page to highlight all of the characteristics
        # good_char_dict is the python dictionary that holds the strings of all the good characteristics of each technology based on the decisions that the user made throughout
        # dict_key is the name of the technology about which the good characteristics are trying to be extracted
        out_string = ""
        if good_char_dict[dict_key] != []: # If the dictionary entry for that key is not empty, then there are some characteristics being stored in there. 
            out_string += "Advantages:"
            for good_char in good_char_dict[dict_key]: # For each good characteristic that has been stored
                out_string += "\n\n+ "+good_char # Add this to the final output string on a new line
        if bad_char_dict[dict_key] != []: # If the dictionary entry for that key is not empty, then there are some characteristics being stored in there. 
            out_string += "\n\nDisadvantages:"
            for bad_char in bad_char_dict[dict_key]:
                out_string += "\n\n- "+ bad_char # Add this to the final output string on a new line
        
        return out_string

    def get_tech_desc(dict_key): # Method to get the strings for the general descriptions of the technologies
        output_string = ""
        match dict_key:
            case "BLE":
                output_string = ("Bluetooth Low Energy (BLE) is the locating technology that utilises low power Bluetooth locators throughout a facility to track a mobile device using its Bluetooth." 
                "When configured correctly it can provide a high accuracy solution (within 3m accuracy) at a lower cost compared to other locating technologies such as Wi-Fi, whose access points may cost several times more.")
            case "Wi-Fi BLE Hybrid":
                output_string=""
            case "Wi-Fi":
                output_string=""
            case "RF":
                output_string=""
            case "DECT":
                output_string=""
            case _:
                output_string="<EXCEPTION REACHED IN get_tech_disc()>"
        return output_string
    
    def print_to_docx(self): # Function to print all the necessary information to a docx file
        print("Get filename...")
        saveDialog = customtkinter.CTkInputDialog(text="Filename:", title="Save as...")
        out_filename = saveDialog.get_input() + ".docx"# Let the user choose the filename for the output file

        print("printing to .docx file...")
        out_doc = Document()
        out_doc.add_picture('img/DocCoverImg.jpg', width=Inches(6)) # Insert a cover image
        out_doc.add_heading(self.gen_desc_dict["Report Title"], 0)
        out_doc.add_page_break() # Start the Text and headings on the next page
        # Introduction
        out_doc.add_heading("Introduction", 1)
        out_doc.add_paragraph(self.gen_desc_dict["Report Intro"])
        # Add recommendations to document
        for i in range(3): # For each of the three recommendations
            rec_heading1 = str(i+1) + ".0 " + self.extract_title(self.top_results[i])
            out_doc.add_heading(rec_heading1,1)
            rec_heading2 = str(i+1) + ".1 " + "Technical Description"
            out_doc.add_heading(rec_heading2, 2)
            out_doc.add_paragraph(self.tech_desc_dict[self.top_results[i]])
            rec_heading3 = str(i+1) + ".2 " + "Characteristics"
            out_doc.add_heading(rec_heading3, 2)
            out_doc.add_paragraph(self.get_char_text(self.good_char, self.bad_char, self.top_results[i]))

        out_doc.save(out_filename)
        messagebox.showinfo("Success", "File printed successfully.")

    def extract_title(self, shortform): # simple function to return the full name of the technology based on the shortform version
        out_string = "EXCEPTION"
        if shortform == "BLE":
            out_string = "Bluetooth Low-Energy (BLE)"
        elif shortform == "Wi-Fi BLE Hybrid":
            out_string = "Wi-Fi Bluetooth Low-Energy Hybrid"
        elif shortform == "Wi-Fi":  
            out_string = "Wi-Fi"
        elif shortform == "RF":
            out_string = "Radio Frequency (RF)"
        elif shortform == "DECT":
            out_string = "Digital Enhanced Cordless Technology (DECT)"
        return out_string

    def save_project(self): # Function to save the project progress to a file
        filepath = filedialog.asksaveasfilename(
        defaultextension=".csv",
        filetypes=[("CSV Files", "*.csv"), ("All Files", "*.*")],
        title="Save File As"
        )
        if filepath: # If the file path exists:
        # save the following to the file in the filepath
            with open(filepath, "w", newline='') as csvfile:
                writer = csv.writer(csvfile) # Create a writer to be able to write data to a csv file
                writer.writerow(["Questions:"]) # Insert a header row to signify the next section
                for questionID in self.question_history:
                    writer.writerow(["Question", questionID]) # Write the questions to the csv file
                writer.writerow(["Choices:"]) # Insert a header row to signify the next section
                for choice in self.choice_history:
                    writer.writerow(["Choice", choice])
                messagebox.showinfo("Success", "File saved successfully.")
        else:
            messagebox.showinfo("ERROR: Could not save file.")
    
    def load_project(self): # Load project from file
        filepath = filedialog.askopenfilename(
            title="Open File",
            filetypes = [("CSV Files", "*.csv"), ("All Files", "*.*")]
        )
        if filepath:
            self.reset_variables()
            with open(filepath, 'r') as csvfile: # Open the file with the specified file path
                reader = csv.reader(csvfile)
                hist_flag = 0 # create a flag to indicate when we are in the question history section
                choice_flag = 0 # create a flag to indictate when we are in the question history section
                for row in reader:
                    print(row)
                    if row[0] == "Questions:": # If we encounter the header, then we are in the questions section
                        hist_flag = 1
                        choice_flag = 0
                    elif row[0] == "Choices:": # If we encounter the header, then we are in the choices section
                        hist_flag = 0
                        choice_flag = 1
                    elif hist_flag == 1:
                        self.question_history.append(row[1])
                    elif choice_flag == 1:
                        self.choice_history.append(int(row[1]))
                    
            self.build_results_page()


    def build_results_page(self):
         ##----------RESULTS PAGE----------##
        print("Change to Results Page")
        try:
            self.WelcomePage.destroy()
        except AttributeError:
            print("Welcome Page does not exist")
        try:
            self.QuestionPage.destroy()
        except AttributeError:
            print("QuestionPage does not exist")

        self.generate_scores(self.choice_history, self.question_history) # use the question and choice history to generate the results.

        self.ResultsPage = customtkinter.CTkFrame(self, fg_color="white", width=self.WIDTH, height=self.HEIGHT) # Creating frame to span the whole page
        self.ResultsPage.grid_columnconfigure((0,5),weight=1)
        self.ResultsPage.grid_columnconfigure((1,2,3,4), weight=8)
        self.ResultsPage.grid_rowconfigure((0,1), weight=4)
        self.ResultsPage.grid_rowconfigure(2, weight=1)

        # Handling the results from the questions
        self.top_results = self.get_top_results(self.scores) # Returns a list of first, second, third keys and scores based on the final scores

                #----WIDGET CONFIGURATIONS----#
        # Top bar
        self.ResultsPage.logo_frame = customtkinter.CTkFrame(self.ResultsPage, fg_color='#183E6F', border_color='#183E6F')
        self.ResultsPage.logo_frame.grid(row=0, column=0, columnspan=5, rowspan=1, sticky="new")
        self.ResultsPage.company_logo_label = customtkinter.CTkLabel(self.ResultsPage.logo_frame, text="Mobile Duress Suitability Software", font=customtkinter.CTkFont(size=20, weight="normal"), text_color="white", fg_color="#183E6F", width=120)  # display image with a CTkLabel
        self.ResultsPage.company_logo_label.grid(row=0, column=0, padx=10, pady=15)

        # Create a frame for each of the buttons and blank spaces in the toolbar

        # Following frames separate the buttons into their own frames to align their spacing
        self.ResultsPage.buttonbar_frame = customtkinter.CTkFrame(self.ResultsPage,width=self.WIDTH-200, height=60, fg_color='#F05A5F', border_width=2, border_color="#F05A5F")
        self.ResultsPage.buttonbar_frame.grid(row=1,column=0, columnspan=5, sticky="nesw")
        #Following grid configurations follow the same as the main page
        self.ResultsPage.buttonbar_frame.grid_columnconfigure((1,2,3), weight=8)
        self.ResultsPage.buttonbar_frame.grid_columnconfigure(0, weight=1)
        self.ResultsPage.buttonbar_frame.grid_columnconfigure(4, weight=1)

        # Following frames separate the buttons into their own frames to align their spacing
        self.ResultsPage.buttonbar_frame0 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.ResultsPage.buttonbar_frame0.grid(row=1,column=0, columnspan=1, sticky="nesw")
        self.ResultsPage.buttonbar_frame1 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.ResultsPage.buttonbar_frame1.grid(row=1,column=1, columnspan=1, sticky="nsew") 
        self.ResultsPage.buttonbar_frame2 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.ResultsPage.buttonbar_frame2.grid(row=1,column=2, columnspan=1, sticky="nesw") 
        self.ResultsPage.buttonbar_frame3 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, width=40, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.ResultsPage.buttonbar_frame3.grid(row=1,column=3, columnspan=1, sticky="nesw")
        self.ResultsPage.buttonbar_frame4 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        self.ResultsPage.buttonbar_frame4.grid(row=1,column=4, columnspan=1, sticky="nesw")
        #self.ResultsPage.buttonbar_frame5 = customtkinter.CTkFrame(self.ResultsPage.buttonbar_frame, height=60, fg_color='#F05A5F', border_color='#F05A5F')
        #self.ResultsPage.buttonbar_frame5.grid(row=1,column=5, columnspan=1, sticky="nesw") 

        # Creating a frame to hold all of the content frames in the white space
        self.content_height=self.HEIGHT-self.ResultsPage.buttonbar_frame.cget("height")-self.ResultsPage.logo_frame.cget("height")-100
        #self.ResultsPage.content_frame = customtkinter.CTkFrame(self.ResultsPage, fg_color="white", width=self.WIDTH, height=self.content_height) 
        self.ResultsPage.content_frame = customtkinter.CTkFrame(self.ResultsPage, fg_color="white") 
        self.ResultsPage.content_frame.grid(row=2, column=1, rowspan=1, columnspan=4, sticky="nsew", padx=20, pady=10)
        self.ResultsPage.content_frame.grid_rowconfigure(0,weight = 1)
        self.ResultsPage.content_frame.grid_rowconfigure(1,weight = 1)
        self.ResultsPage.content_frame.grid_rowconfigure(2,weight = 5)
        self.ResultsPage.content_frame.grid_rowconfigure(3,weight = 5)
        self.ResultsPage.content_frame.grid_columnconfigure(0, weight = 1)
        #self.ResultsPage.content_frame.grid_columnconfigure(1,weight=4)

        # Creating buttons
        #self.ResultsPage.save_project_button = customtkinter.CTkButton(self.ResultsPage.buttonbar_frame1, command=self.home_button_event, text='Save Current Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        #self.ResultsPage.save_project_button.pack(padx=5, pady=10)
        self.ResultsPage.load_button = customtkinter.CTkButton(self.ResultsPage.buttonbar_frame1, command=self.load_project, text='Load Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        self.ResultsPage.load_button.pack(padx=5, pady=10)
        self.ResultsPage.new_project_button = customtkinter.CTkButton(self.ResultsPage.buttonbar_frame3, command=self.changeToQuestion1, text='New Project', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        self.ResultsPage.new_project_button.pack(padx=5, pady=10)
        self.ResultsPage.home_button = customtkinter.CTkButton(self.ResultsPage.buttonbar_frame2, command=self.home_button_event, text='Home', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        self.ResultsPage.home_button.pack(padx=5, pady=10)

        # Contents on page
        # Labels
        self.ResultsPage.content_frame.heading = customtkinter.CTkLabel(self.ResultsPage.content_frame, text="Results", font=customtkinter.CTkFont(size=24, weight="bold"), text_color="#183E6F")
        self.ResultsPage.content_frame.heading.grid(row=0, column=0, padx=20, pady=10)
        self.ResultsPage.content_frame.description = customtkinter.CTkLabel(
            self.ResultsPage.content_frame,
            text=self.gen_desc_dict["Results Desc"],
            font=customtkinter.CTkFont(size=14, weight="normal"),
            text_color="#183E6F",
            wraplength=1200)
        self.ResultsPage.content_frame.description.grid(row=1, column=0, padx=20, pady=10)

        # Tab views
        self.ResultsPage.results_tab = customtkinter.CTkTabview(
            self.ResultsPage.content_frame, width=self.WIDTH, 
            fg_color="#a0ddff", corner_radius=6, 
            segmented_button_fg_color="#F05A5F", 
            segmented_button_selected_color="#183E6F",
            segmented_button_unselected_color="#F05A5F",
            segmented_button_selected_hover_color="#183E6F",
            segmented_button_unselected_hover_color="#183E6F")
        
        self.ResultsPage.results_tab.grid(row=2,column=0, padx=20, pady=20)
        self.ResultsPage.results_tab.add("1st Recommendation")
        self.ResultsPage.results_tab.add("2nd Recommendation")
        self.ResultsPage.results_tab.add("3rd Recommendation")

        # Creating the first tab
        # Scrollable frame
        self.ResultsPage.scrollframe1 = customtkinter.CTkScrollableFrame(self.ResultsPage.results_tab.tab("1st Recommendation"), width=500, height=300, fg_color="#a0ddff")
        self.ResultsPage.scrollframe1.grid(row=1, column=1, sticky="nsew")

        self.ResultsPage.results_tab.tab("1st Recommendation").grid_columnconfigure(0,weight=1)
        self.ResultsPage.results_tab.tab("1st Recommendation").grid_columnconfigure(1,weight=40)
        self.ResultsPage.results_tab.tab("1st Recommendation").grid_rowconfigure(0, weight=1) # Weights for the headings
        self.ResultsPage.results_tab.tab("1st Recommendation").grid_rowconfigure(1, weight=8) # Weights for the body text

        
        self.ResultsPage.scrollframe1.grid_rowconfigure((0,2), weight=1) # Weights for the headings
        self.ResultsPage.scrollframe1.grid_rowconfigure((1,3), weight=8) # Weights for the body text

        #Putting widgets inside the scrollable frame inside the first tab
        self.ResultsPage.results_tab.number_tag = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("1st Recommendation"), text="1", font=customtkinter.CTkFont(size=14, weight="bold"), text_color="white", corner_radius=500, fg_color="#F05A5F", width=30, height=30)
        self.ResultsPage.results_tab.number_tag.grid(row=0, column=0, padx=5, pady=5)
        self.ResultsPage.results_tab.results_techheading1 = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("1st Recommendation"), text=self.extract_title(self.top_results[0]), font=customtkinter.CTkFont(size=24, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_techheading1.grid(row=0, column=1, sticky="nesw", padx=30, pady=10)
        # self.ResultsPage.results_tab.results_description1 = customtkinter.CTkLabel(self.ResultsPage.scrollframe1, text=self.tech_desc_dict[self.top_results[0]], font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", wraplength=1200)
        # self.ResultsPage.results_tab.results_description1.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_charheading1 = customtkinter.CTkLabel(self.ResultsPage.scrollframe1, text="Characteristics", font=customtkinter.CTkFont(size=20, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_charheading1.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_characteristics = customtkinter.CTkLabel(self.ResultsPage.scrollframe1, text=self.get_char_text(self.good_char, self.bad_char, self.top_results[0]), font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", fg_color = "#a0ddff", wraplength=1200)
        self.ResultsPage.results_tab.results_characteristics.grid(row=2, column=0, sticky="nesw", padx=30, pady=10)

        # Creating the second tab
        # Scrollable frame
        self.ResultsPage.scrollframe2 = customtkinter.CTkScrollableFrame(self.ResultsPage.results_tab.tab("2nd Recommendation"), width=500, height=300, fg_color="#a0ddff")
        self.ResultsPage.scrollframe2.grid(row=1, column=1, sticky="nsew")

        self.ResultsPage.results_tab.tab("2nd Recommendation").grid_columnconfigure(0,weight=1)
        self.ResultsPage.results_tab.tab("2nd Recommendation").grid_columnconfigure(1,weight=40)
        self.ResultsPage.results_tab.tab("2nd Recommendation").grid_rowconfigure(0, weight=1) # Weights for the headings
        self.ResultsPage.results_tab.tab("2nd Recommendation").grid_rowconfigure(1, weight=8) # Weights for the body text

        
        self.ResultsPage.scrollframe2.grid_rowconfigure((0,2), weight=1) # Weights for the headings
        self.ResultsPage.scrollframe2.grid_rowconfigure((1,3), weight=8) # Weights for the body text

        #Putting widgets inside the scrollable frame inside the first tab
        self.ResultsPage.results_tab.number_tag = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("2nd Recommendation"), text="2", font=customtkinter.CTkFont(size=14, weight="bold"), text_color="white", corner_radius=500, fg_color="#F05A5F", width=30, height=30)
        self.ResultsPage.results_tab.number_tag.grid(row=0, column=0, padx=5, pady=5)
        self.ResultsPage.results_tab.results_techheading2 = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("2nd Recommendation"), text=self.extract_title(self.top_results[1]), font=customtkinter.CTkFont(size=24, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_techheading2.grid(row=0, column=1, sticky="nesw", padx=30, pady=10)
        # self.ResultsPage.results_tab.results_description2 = customtkinter.CTkLabel(self.ResultsPage.scrollframe2, text=self.tech_desc_dict[self.top_results[1]], font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", wraplength=1200)
        # self.ResultsPage.results_tab.results_description2.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_charheading2 = customtkinter.CTkLabel(self.ResultsPage.scrollframe2, text="Characteristics", font=customtkinter.CTkFont(size=20, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_charheading2.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_characteristics = customtkinter.CTkLabel(self.ResultsPage.scrollframe2, text=self.get_char_text(self.good_char, self.bad_char, self.top_results[1]), font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", fg_color = "#a0ddff", wraplength=1200)
        self.ResultsPage.results_tab.results_characteristics.grid(row=2, column=0, sticky="nesw", padx=30, pady=10)

        # Creating the third tab
        # Scrollable frame
        self.ResultsPage.scrollframe3 = customtkinter.CTkScrollableFrame(self.ResultsPage.results_tab.tab("3rd Recommendation"), width=500, height=300, fg_color="#a0ddff")
        self.ResultsPage.scrollframe3.grid(row=1, column=1, sticky="nsew")

        self.ResultsPage.results_tab.tab("3rd Recommendation").grid_columnconfigure(0,weight=1)
        self.ResultsPage.results_tab.tab("3rd Recommendation").grid_columnconfigure(1,weight=40)
        self.ResultsPage.results_tab.tab("3rd Recommendation").grid_rowconfigure(0, weight=1) # Weights for the headings
        self.ResultsPage.results_tab.tab("3rd Recommendation").grid_rowconfigure(1, weight=8) # Weights for the body text

        
        self.ResultsPage.scrollframe3.grid_rowconfigure((0,2), weight=1) # Weights for the headings
        self.ResultsPage.scrollframe3.grid_rowconfigure((1,3), weight=8) # Weights for the body text

        #Putting widgets inside the scrollable frame inside the first tab
        self.ResultsPage.results_tab.number_tag = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("3rd Recommendation"), text="3", font=customtkinter.CTkFont(size=14, weight="bold"), text_color="white", corner_radius=500, fg_color="#F05A5F", width=30, height=30)
        self.ResultsPage.results_tab.number_tag.grid(row=0, column=0, padx=5, pady=5)
        self.ResultsPage.results_tab.results_techheading3 = customtkinter.CTkLabel(self.ResultsPage.results_tab.tab("3rd Recommendation"), text=self.extract_title(self.top_results[2]), font=customtkinter.CTkFont(size=24, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_techheading3.grid(row=0, column=1, sticky="nesw", padx=30, pady=10)
        # self.ResultsPage.results_tab.results_description3 = customtkinter.CTkLabel(self.ResultsPage.scrollframe3, text=self.tech_desc_dict[self.top_results[2]], font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", wraplength=1200)
        # self.ResultsPage.results_tab.results_description3.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_charheading3 = customtkinter.CTkLabel(self.ResultsPage.scrollframe3, text="Characteristics", font=customtkinter.CTkFont(size=20, weight="bold"), text_color="#183E6F", anchor="w")
        self.ResultsPage.results_tab.results_charheading3.grid(row=1, column=0, sticky="nesw", padx=30, pady=10)
        self.ResultsPage.results_tab.results_characteristics = customtkinter.CTkLabel(self.ResultsPage.scrollframe3, text=self.get_char_text(self.good_char, self.bad_char, self.top_results[2]), font=customtkinter.CTkFont(size=14, weight="normal"), text_color="#183E6F", anchor="w", justify="left", fg_color = "#a0ddff", wraplength=1200)
        self.ResultsPage.results_tab.results_characteristics.grid(row=2, column=0, sticky="nesw", padx=30, pady=10)

        # Creating the frame to hold the save and print to docx. buttons
        self.ResultsPage.prevnext_frame = customtkinter.CTkFrame(self.ResultsPage.content_frame, fg_color="white", height=40)
        self.ResultsPage.prevnext_frame.grid(row=3, column=0, sticky="ew", pady=10, padx=10)
        self.ResultsPage.prevnext_frame.grid_columnconfigure(0,weight=15)
        self.ResultsPage.prevnext_frame.grid_columnconfigure((1,2),weight=1)
        self.ResultsPage.prevnext_frame.grid_rowconfigure(0,weight=1)

        # Creating Previous and next buttons at the bottom of the page
        self.ResultsPage.save_button = customtkinter.CTkButton(self.ResultsPage.prevnext_frame, command=self.save_project, text='Save', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        self.ResultsPage.save_button.grid(row=2, column=2)
        self.ResultsPage.print_button = customtkinter.CTkButton(self.ResultsPage.prevnext_frame, command=self.print_to_docx, text='Print to .docx', fg_color="#183E6F", border_color="#183E6F", font=customtkinter.CTkFont(size=14, weight="bold"), height=40, hover_color="#e26d5c")
        self.ResultsPage.print_button.grid(row=2, column=3)
        self.ResultsPage.print_button.configure(state="enabled") # Can't go to the next question until we put in a radio option

        self.ResultsPage.pack(fill=customtkinter.BOTH)


    
if __name__ == "__main__":
    app = App()
    app.mainloop()